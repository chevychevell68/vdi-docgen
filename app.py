import os
import json
import base64
import zipfile
import uuid
import datetime as dt
from pathlib import Path
from flask import (
    Flask, request, render_template, redirect, url_for, flash,
    send_file, current_app
)
from urllib.request import Request, urlopen
from urllib.error import HTTPError, URLError

# Optional dependency for PDG parsing (python-docx)
try:
    from docx import Document  # type: ignore
except Exception:
    Document = None  # We'll error nicely if upload used without this pkg

app = Flask(__name__, template_folder="templates", static_folder="static")
app.config["SECRET_KEY"] = os.getenv("SECRET_KEY", "dev-secret-change-me")
# Application version (bump as needed)
APP_VERSION = "2025.10.14-v6"

# ---------- Jinja helpers ----------
def _fmt_bool(v):
    return "Yes" if str(v).lower() in {"1", "true", "yes", "on"} else "No"

def _none_to_empty(v):
    return "" if v is None else v

def _lines(s):
    if not s:
        return []
    return [line.rstrip() for line in str(s).splitlines()]

app.jinja_env.filters["fmt_bool"] = _fmt_bool
app.jinja_env.filters["none_to_empty"] = _none_to_empty
app.jinja_env.filters["lines"] = _lines

# -------------------- Paths / Env --------------------
OUTPUT_DIR   = Path(os.getenv("OUTPUT_DIR", "output"))
DELIV_DIR    = OUTPUT_DIR / "deliverables"
DOCX_DIR     = Path(os.getenv("DOCX_DIR", "output-docx"))
EXPORTS_DIR  = OUTPUT_DIR / "exports"
SUBMIT_DIR   = Path(os.getenv("SUBMIT_DIR", "submissions"))  # repo-local for Codespaces

for d in (OUTPUT_DIR, DELIV_DIR, DOCX_DIR, EXPORTS_DIR, SUBMIT_DIR):
    d.mkdir(parents=True, exist_ok=True)

# Known/expected DOCX deliverables (used for ctx + whitelisting)
KNOWN_DOCS = [
    "SOW.docx", "HLD.docx", "LLD.docx", "Runbook.docx",
    "PDG.docx", "LOE-WBS.docx", "ATP.docx", "Adoption_Plan.docx",
    # Historical names we may also emit
    "LOE.docx", "WBS.docx"
]

GITHUB_TOKEN   = os.getenv("GITHUB_TOKEN", "").strip()
DEFAULT_REPO   = os.getenv("DEFAULT_REPO", "").strip()  # e.g., "owner/repo"
DEFAULT_BRANCH = os.getenv("DEFAULT_BRANCH", "main").strip()

@app.context_processor
def inject_helpers():
    def has_endpoint(name: str) -> bool:
        try:
            return name in current_app.view_functions
        except Exception:
            return False
    # expose UTC ‘now’ so templates can call now.strftime(...)
    return dict(
        has_endpoint=has_endpoint,
        now=dt.datetime.utcnow(),
        app_version=APP_VERSION,
        output_dir=str(OUTPUT_DIR),
        docx_dir=str(DOCX_DIR),
        submit_dir=str(SUBMIT_DIR),
    )

# -------------------- GitHub helpers (optional) --------------------
def github_get_file_sha(repo_owner: str, repo_name: str, path: str, branch: str):
    if not (GITHUB_TOKEN and repo_owner and repo_name and path and branch):
        return None
    url = f"https://api.github.com/repos/{repo_owner}/{repo_name}/contents/{path}?ref={branch}"
    req = Request(url, headers={"Authorization": f"token {GITHUB_TOKEN}", "Accept": "application/vnd.github+json"})
    try:
        with urlopen(req) as resp:
            data = json.loads(resp.read().decode("utf-8"))
            return data.get("sha")
    except HTTPError as e:
        if e.code == 404:
            return None
        return None
    except URLError:
        return None
    except Exception:
        return None

def github_upsert_file(repo_owner: str, repo_name: str, branch: str, path: str, content_bytes: bytes, commit_msg: str) -> bool:
    if not (GITHUB_TOKEN and repo_owner and repo_name and branch and path):
        return False
    url = f"https://api.github.com/repos/{repo_owner}/{repo_name}/contents/{path}"
    body = {
        "message": commit_msg,
        "content": base64.b64encode(content_bytes).decode("utf-8"),
        "branch": branch,
    }
    existing_sha = github_get_file_sha(repo_owner, repo_name, path, branch)
    if existing_sha:
        body["sha"] = existing_sha

    data = json.dumps(body).encode("utf-8")
    req = Request(
        url,
        data=data,
        method="PUT",
        headers={
            "Authorization": f"token {GITHUB_TOKEN}",
            "Accept": "application/vnd.github+json",
            "Content-Type": "application/json",
        },
    )
    try:
        with urlopen(req) as resp:
            _ = resp.read()
        return True
    except Exception:
        return False

def push_export_to_github(owner_repo: str, branch: str, rel_path: str, content: bytes, commit_msg: str) -> bool:
    if not owner_repo:
        return False
    try:
        owner, repo = owner_repo.split("/", 1)
    except ValueError:
        return False
    return github_upsert_file(owner, repo, branch, rel_path, content, commit_msg)

# -------------------- Helpers --------------------
def _now_str() -> str:
    return dt.datetime.now().strftime("%Y-%m-%d %H:%M")

def _uuid() -> str:
    return uuid.uuid4().hex

def _save_submission_local(payload: dict, submit_id: str) -> Path:
    payload["_id"] = submit_id
    payload["_saved_at"] = _now_str()
    p = SUBMIT_DIR / f"{submit_id}.json"
    p.write_text(json.dumps(payload, indent=2), encoding="utf-8")
    return p

def _save_submission(payload: dict, submit_id: str | None = None) -> str:
    submit_id = submit_id or payload.get("_id") or _uuid()
    _save_submission_local(payload, submit_id)
    # (Optional) push to GitHub to mirror your old Render flow
    if GITHUB_TOKEN and DEFAULT_REPO:
        try:
            rel = f"submissions/{submit_id}.json"
            push_export_to_github(DEFAULT_REPO, DEFAULT_BRANCH, rel, json.dumps(payload, indent=2).encode("utf-8"),
                                  f"Presales submission {submit_id}")
        except Exception:
            pass
    return submit_id

def _load_submission(submit_id: str):
    p = SUBMIT_DIR / f"{submit_id}.json"
    if not p.exists():
        return None
    try:
        return json.loads(p.read_text(encoding="utf-8"))
    except Exception:
        return None

def _map_requested_docx(requested_labels, base_dir: Path):
    label_to_filename = {
        "SOW": "SOW.docx",
        "HLD": "HLD.docx",
        "LLD": "LLD.docx",
        "Runbook": "Runbook.docx",
        "PDG": "PDG.docx",
        "LOE/WBS": "LOE-WBS.docx",
        "ATP": "ATP.docx",
        "Adoption Plan": "Adoption_Plan.docx",
    }
    requested_fns = [label_to_filename.get(lbl) for lbl in requested_labels if lbl in label_to_filename]
    available, missing = [], []
    for fn in requested_fns:
        if not fn:
            continue
        p = base_dir / fn
        (available if p.exists() else missing).append(fn)
    return available, missing

def _build_submission_zip(submit_id: str, filenames: list[str]) -> Path:
    export_root = EXPORTS_DIR / "submissions"
    export_root.mkdir(parents=True, exist_ok=True)
    zip_path = export_root / f"presales-{submit_id}.zip"
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as z:
        for fn in filenames:
            fp = DOCX_DIR / fn
            if fp.exists():
                z.write(fp, arcname=fn)
    return zip_path

# ---- ctx helpers for template buttons ----
def _build_ctx(submit_id: str, obj: dict | None) -> dict:
    """Construct a minimal ctx expected by the template.
    We include any KNOWN_DOCS present in DOCX_DIR, and prefer the originally requested docs if available.
    """
    # Prefer requested doc names when available
    requested = (obj or {}).get("docs_requested_list") or []
    # Map requested labels to filenames (like _map_requested_docx)
    label_to_filename = {
        "SOW": "SOW.docx", "HLD": "HLD.docx", "LLD": "LLD.docx", "Runbook": "Runbook.docx",
        "PDG": "PDG.docx", "LOE/WBS": "LOE-WBS.docx", "ATP": "ATP.docx", "Adoption Plan": "Adoption_Plan.docx",
    }
    requested_files = [label_to_filename.get(lbl) for lbl in requested if lbl in label_to_filename]
    present = {p.name for p in DOCX_DIR.glob("*.docx")}
    # If there are requested files, intersect with present; else include any known present
    candidates = [fn for fn in requested_files if fn in present] if requested_files else [fn for fn in KNOWN_DOCS if fn in present]
    deliverables = [{"filename": n, "title": n.rsplit(".", 1)[0]} for n in candidates]
    return {
        "sid": submit_id,
        "deliverables": deliverables,
        "docx_dir": str(DOCX_DIR),
    }

# -------------------- Routes --------------------
@app.get("/")
def index():
    return render_template("index.html", storage_dir=str(SUBMIT_DIR), app_version=APP_VERSION)

@app.get("/predeploy")
def predeploy():
    return render_template("predeploy.html")

@app.get("/checklist")
def checklist():
    return render_template("checklist.html")

@app.get("/questionnaire")
def questionnaire():
    return render_template("questionnaire.md")

# ---- PDG (unchanged quick MD render form) ----
@app.get("/pdg")
def pdg_form():
    return render_template("pdg_form.html")

@app.post("/pdg/submit")
def pdg_submit():
    data = {k: v for k, v in request.form.items()}
    data["__submitted_at__"] = _now_str()
    md = render_template("pdg.md.j2", **data, data=data, now=dt.datetime.utcnow())
    (DELIV_DIR / "PDG.md").write_text(md, encoding="utf-8")
    html = render_template("pdg_results.html", **data, rendered_md=md)
    return html

# ---- Downloads (MD/DOCX) ----
@app.get("/download/md/<name>")
def download_md(name: str):
    path = (DELIV_DIR / name).resolve()
    if not path.exists():
        flash("File not found.", "error")
        return redirect(url_for("history"))
    return send_file(path, as_attachment=True, download_name=name)

@app.get("/download/docx/<name>")
def download_docx(name: str):
    path = (DOCX_DIR / name).resolve()
    if not path.exists():
        flash("File not found.", "error")
        return redirect(url_for("history"))
    return send_file(path, as_attachment=True, download_name=name)

# New: safe individual deliverable download used by the template helper
@app.get("/download/<submit_id>/<path:filename>")
def download_deliverable(submit_id: str, filename: str):
    # Allow only known deliverables that actually exist in DOCX_DIR
    if filename not in KNOWN_DOCS and not (DOCX_DIR / filename).exists():
        flash("File not found.", "error")
        return redirect(url_for("presales_view", submit_id=submit_id))
    path = (DOCX_DIR / filename).resolve()
    if not path.exists():
        flash("File not found.", "error")
        return redirect(url_for("presales_view", submit_id=submit_id))
    return send_file(path, as_attachment=True, download_name=filename)

# ---- Exports (general) ----
@app.get("/exports")
def exports():
    zips = sorted(EXPORTS_DIR.glob("*.zip"))
    return render_template("exports.html", zips=[f.name for f in zips])

# ---- Presales (form / submit / view / zip / edit) ----
@app.get("/presales")
def presales():
    return render_template("presales_form.html", form=request.form)

@app.post("/presales/submit")
def presales_submit():
    # Capture full form data
    data = request.form.to_dict(flat=True)
    data["__submitted_at__"] = _now_str()

    # Preserve multi-selects / checkboxes as lists
    data["docs_requested_list"] = request.form.getlist("docs_requested")
    data["profile_mgmt_list"] = request.form.getlist("profile_mgmt")
    data["virtual_apps_list"] = request.form.getlist("virtual_apps")
    data["training_required_list"] = request.form.getlist("training_required")

    # Regions: capture all region flags & pct values that exist
    region_keys = [
        "US","US_HI","US_AK","CAN","LATAM","EMEA","APAC","INDIA","ANZ","OTHER"
    ]
    for rk in region_keys:
        data[f"region_ck_{rk}"]  = request.form.get(f"region_ck_{rk}")
        data[f"region_pct_{rk}"] = request.form.get(f"region_pct_{rk}")

    # Save JSON snapshot
    submit_id = _save_submission(data, data.get("_id"))

    # Optional MD snapshot for audit/troubleshooting (not offered for download here)
    md = render_template("presales_export.md", data=data, **data, now=dt.datetime.utcnow())
    (DELIV_DIR / f"intake-{submit_id}.md").write_text(md, encoding="utf-8")

    # Redirect to the results page
    return redirect(url_for("presales_view", submit_id=submit_id))

@app.get("/presales/view/<submit_id>")
def presales_view(submit_id):
    obj = _load_submission(submit_id)
    if not obj:
        flash("Submission not found.", "error")
        return redirect(url_for("history"))

    requested = obj.get("docs_requested_list") or []
    docx_available, docx_missing = _map_requested_docx(requested, DOCX_DIR)

    # Build a ctx object for the template so individual buttons & links render
    ctx = _build_ctx(submit_id, obj)

    return render_template(
        "presales_submitted.html",
        what="Presales",
        data=obj,
        submit_id=submit_id,
        docx_available=docx_available,
        docx_missing=docx_missing,
        ctx=ctx,  # <-- critical for the template buttons
    )

@app.get("/presales/zip/<submit_id>")
def presales_zip(submit_id):
    obj = _load_submission(submit_id)
    if not obj:
        flash("Submission not found.", "error")
        return redirect(url_for("history"))
    requested = obj.get("docs_requested_list") or []
    have, _missing = _map_requested_docx(requested, DOCX_DIR)
    # If nothing requested is available, zip up any known docs present
    if not have:
        have = [n for n in KNOWN_DOCS if (DOCX_DIR / n).exists()]
    if not have:
        flash("No requested Word documents available to zip.", "warning")
        return redirect(url_for("presales_view", submit_id=submit_id))
    zip_path = _build_submission_zip(submit_id, have)
    return send_file(zip_path, as_attachment=True, download_name=zip_path.name)

@app.get("/presales/edit/<submit_id>")
def presales_edit(submit_id):
    obj = _load_submission(submit_id)
    if not obj:
        flash("Submission not found.", "error")
        return redirect(url_for("history"))
    class _F(dict):
        __getattr__ = dict.get
    f = _F(obj.copy())
    return render_template("presales_form.html", form=f, submit_id=submit_id)

# ---- PDG Upload & Merge (.docx -> presales JSON) ----
# Single endpoint handles GET (form) and POST (upload)
@app.route("/presales/upload-pdg/<submit_id>", methods=["GET", "POST"])
def presales_upload_pdg(submit_id):
    obj = _load_submission(submit_id)
    if not obj:
        flash("Submission not found.", "error")
        return redirect(url_for("history"))

    if request.method == "GET":
        return render_template("pdg_upload.html", submit_id=submit_id, data=obj)

    # POST (upload)
    if Document is None:
        flash("python-docx is not installed on the server.", "error")
        return redirect(url_for("presales_view", submit_id=submit_id))

    f = request.files.get("pdg_file")
    if not f or not f.filename.lower().endswith(".docx"):
        flash("Please upload a .docx PDG file.", "error")
        return redirect(url_for("presales_upload_pdg", submit_id=submit_id))

    # Save uploaded PDG under submissions/<sid>/uploads
    updir = SUBMIT_DIR / submit_id / "uploads"
    updir.mkdir(parents=True, exist_ok=True)
    stamp = dt.datetime.utcnow().strftime("%Y%m%d-%H%M%S")
    saved = updir / f"PDG-{stamp}.docx"
    f.save(saved)

    # Parse PDG
    try:
        parsed = _parse_pdg_docx(saved)
    except Exception as e:
        flash(f"Failed to parse PDG: {e}", "error")
        return redirect(url_for("presales_view", submit_id=submit_id))

    # Merge into presales data (fill blanks by default)
    obj, n = _merge_pdg_into_presales(parsed, obj)
    obj.setdefault("_pdg_uploads", [])
    obj["_pdg_uploads"].append({
        "path": str(saved),
        "merged": n,
        "at": dt.datetime.utcnow().isoformat()
    })
    _save_submission(obj)

    if n == 0:
        flash("PDG uploaded. No new fields detected to merge.", "warning")
    else:
        flash(f"PDG uploaded and merged {n} field(s) into presales data.", "success")

    return redirect(url_for("presales_view", submit_id=submit_id))

# ---- PDG parsing + mapping ----
def _parse_pdg_docx(path: Path) -> dict:
    import re
    doc = Document(str(path))
    out = {}

    def norm_key(k: str) -> str:
        k = (k or "").strip().lower()
        k = re.sub(r"[^a-z0-9]+", "_", k).strip("_")
        return k

    # Tables: treat 2-col rows as key:value
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) >= 2:
                key, val = cells[0], cells[1]
                if key and val:
                    out[norm_key(key)] = val

    # Paragraphs: "Key: Value"
    for p in doc.paragraphs:
        txt = (p.text or "").strip()
        if ":" in txt:
            key, val = txt.split(":", 1)
            key = key.strip()
            val = val.strip()
            if key and val:
                out[norm_key(key)] = val

    return out

# PDG normalized keys -> presales JSON fields
_PDG_TO_PRESALES = {
    # Customer info
    "company": "company_name",
    "company_name": "company_name",
    "customer_name": "customer_name",
    "primary_contact": "customer_name",
    "sf_opportunity_name": "sf_opportunity_name",
    "salesforce_opportunity_url": "sf_opportunity_url",
    "salesforce_opportunity": "sf_opportunity_name",
    # Users & scope
    "total_users": "total_users",
    "concurrent_users": "concurrent_users",
    "num_datacenters": "num_datacenters",
    "datacenters_regions_where": "datacenters_detail",
    "deployment_type": "deployment_type",
    "main_use_case": "main_use_cases",
    # Identity / access
    "mfa_required": "mfa_required",
    "remote_access_required": "remote_access",
    "idp": "idp_provider",
    # GPUs
    "gpu_required": "gpu_required",
    "gpu_users": "gpu_users",
    "gpu_vram_per_user_gb": "gpu_vram_per_user",
    "gpu_use_case": "gpu_use_case",
    # Platform
    "platform": "platform",
    "storage_type": "storage_type",
    "load_balancer": "load_balancer",
    # Logistics
    "delivery_model": "delivery_model",
    "start_date": "start_date",
    "timeline_deadlines": "timeline",
}

def _merge_pdg_into_presales(pdg_fields: dict, obj: dict) -> tuple[dict, int]:
    updated = 0
    for k_norm, val in pdg_fields.items():
        target = _PDG_TO_PRESALES.get(k_norm)
        if not target:
            continue
        prev = obj.get(target)
        if not prev or str(prev).strip() == "":  # fill blanks only
            obj[target] = val
            updated += 1
    return obj, updated

# ---- History (reads repo-local submissions/) ----
@app.get("/history")
def history():
    rows = []
    SUBMIT_DIR.mkdir(parents=True, exist_ok=True)
    for p in sorted(SUBMIT_DIR.glob("*.json")):
        try:
            obj = json.loads(p.read_text(encoding="utf-8"))
        except Exception:
            continue
        rows.append({
            "id": obj.get("_id"),
            "company_name": (obj.get("company_name") or "—").strip(),
            "sf_op_name": (obj.get("sf_opportunity_name") or "—").strip(),
            "sf_url": (obj.get("sf_opportunity_url") or "").strip(),
            "submitted_at": (obj.get("__submitted_at__") or obj.get("_saved_at") or ""),
        })
    rows.sort(key=lambda r: r.get("submitted_at",""), reverse=True)
    return render_template("history.html", rows=rows)

# ---- Simple 404/500 ----
@app.errorhandler(404)
def not_found(_e):
    return (
        '<!doctype html><html><head><meta charset="utf-8"><title>Not Found</title>'
        '<meta name="viewport" content="width=device-width, initial-scale=1"></head>'
        '<body style="font-family:system-ui,-apple-system,Segoe UI,Roboto,Arial;padding:2rem">'
        "<h1>404 — Not Found</h1>"
        f'<p><a href="{url_for("index")}">Go to Home</a></p>'
        "</body></html>"
    ), 404

@app.errorhandler(500)
def internal_error(_e):
    return (
        '<!doctype html><html><head><meta charset="utf-8"><title>Server Error</title>'
        '<meta name="viewport" content="width=device-width, initial-scale=1"></head>'
        '<body style="font-family:system-ui,-apple-system,Segoe UI,Roboto,Arial;padding:2rem">'
        "<h1>500 — Internal Server Error</h1>"
        f'<p><a href="{url_for("index")}">Go to Home</a></p>'
        "</body></html>"
    ), 500

# ---- Optional blueprint support (safe no-op if not present) ----
try:
    from presales import presales_bp
    app.register_blueprint(presales_bp)
except Exception:
    pass

if __name__ == "__main__":
    port = int(os.getenv("PORT", "5000"))
    app.run(host="0.0.0.0", port=port, debug=os.getenv("DEBUG", "1") == "1")
