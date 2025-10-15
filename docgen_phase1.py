# docgen_phase1.py (additive module — no changes to app.py required)
import re, json, zipfile
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Tuple

try:
    from docx import Document
    from docx.shared import Pt
except Exception:
    Document = None  # App remains functional even without python-docx

def _ts() -> str:
    return datetime.now().strftime("%Y%m%d-%H%M")

def _initials(text: str, max_len=5) -> str:
    if not text: return ""
    ups = "".join([c for c in text if c.isalpha() and c.isupper()])[:max_len]
    if ups: return ups
    parts = re.findall(r"[A-Za-z0-9]+", text)
    return "".join(p[0] for p in parts)[:max_len].upper()

def abbr_customer(name: str) -> str:
    if not name: return "CUST"
    if re.search(r"\bPACAF\b|\bPacific\s+Air\s+Forces\b", name, re.I): return "PACAF"
    if re.search(r"\bDiamondback\s+Energy\b", name, re.I): return "DBE"
    return _initials(name, 5)

def abbr_project(text: str) -> str:
    if not text: return "PRJ"
    t = text.strip()
    toks, words = [], re.findall(r"[A-Za-z0-9]+", t)
    for w in words:
        wl = w.lower()
        if wl.startswith("horizon"): toks.append("HZN")
        elif wl in ("vcf","vcloud","cloudfoundation","cloud-foundation"): toks.append("VCF")
        elif wl in ("deploy","deployment","impl","implementation"): toks.append("IMPL")
        elif wl == "pod": toks.append("POD")
        elif re.fullmatch(r"pod\\d+", wl): toks.append(w.upper())
        else:
            if len(w) <= 3 or w.isdigit(): toks.append(w.upper())
    base = "".join(toks) or _initials(t, 8)
    return re.sub(r"[^A-Z0-9]", "", base)[:10] or "PRJ"

def _normal_font(doc):
    style = doc.styles["Normal"]
    f = style.font
    f.name, f.size = "Calibri", Pt(11)

def _p(doc, text):
    p = doc.add_paragraph(text or "")
    for r in p.runs: r.bold = False

def _h(doc, text):
    p = doc.add_paragraph(); r = p.add_run(text or ""); r.bold = False
    doc.add_paragraph("")

def _kv(doc, rows: List[Tuple[str,str]]):
    if not rows: return
    t = doc.add_table(rows=len(rows), cols=2); t.autofit = True
    for i,(k,v) in enumerate(rows): t.cell(i,0).text, t.cell(i,1).text = k or "", v or ""
    doc.add_paragraph("")

def _safe(ctx: Dict, key: str, default=""):
    v = ctx.get(key, default)
    return default if v is None else v

def generate_all(ctx: Dict, outdir: Path) -> Dict[str, Path]:
    """
    Create SOW/HLD/LOE docs into outdir using presales data in ctx.
    Returns map: {'SOW': path, 'HLD': path, 'LOE': path, 'ZIP': path}
    No exceptions bubble out: if python-docx missing, returns {}.
    """
    if Document is None:
        return {}
    outdir.mkdir(parents=True, exist_ok=True)
    cust_code, proj_code = abbr_customer(_safe(ctx,"customer_name","Customer")), abbr_project(_safe(ctx,"project_name","Project"))
    ts = _ts()

    # --- SOW (brief but compliant to your style)
    sow = outdir / f"{cust_code}_{proj_code}_SOW_{ts}.docx"
    doc = Document(); _normal_font(doc)
    _h(doc, f"Statement of Work\n{_safe(ctx,'project_name','Project')}\n{_safe(ctx,'customer_name','Customer')}")
    _h(doc, "Introduction")
    _p(doc, "This Statement of Work describes the services to design and deploy the Omnissa Horizon environment. Work is executed under the EPDIO framework (Engage, Plan, Design, Implement, Operate).")
    _h(doc, "Scope of Work")
    _p(doc, "\n".join([
        f"Target users: {_safe(ctx,'total_users', _safe(ctx,'user_count','(from presales)'))}",
        f"Regions: {_safe(ctx,'regions','(from presales)')}",
        f"GPU: {_safe(ctx,'gpu_required','No')}",
        f"External access: {_safe(ctx,'remote_access', _safe(ctx,'external_access','(from presales)'))}",
        f"VCF domains: {_safe(ctx,'vcf_domains','(from presales)')}",
    ]))
    _h(doc, "EPDIO Phases"); _p(doc, "Engage / Plan / Design / Implement / Operate — see HLD & LLD for details.")
    _h(doc, "Deliverables"); _p(doc, "HLD, LLD, Implemented Horizon, ATP results, As-Built, KT")
    _h(doc, "Acceptance Criteria"); _p(doc, "ATP executed successfully; configuration aligns with approved LLD.")
    _kv(doc, [("Customer", _safe(ctx,"customer_name")),("Project", _safe(ctx,"project_name")),("User Count", str(_safe(ctx,'user_count', '')))])
    doc.save(sow)

    # --- HLD (basic conceptual/logical/physical)
    hld = outdir / f"{cust_code}_{proj_code}_HLD_{ts}.docx"
    doc = Document(); _normal_font(doc)
    _h(doc, f"High-Level Design\n{_safe(ctx,'project_name')}\n{_safe(ctx,'customer_name')}")
    _h(doc, "Conceptual"); _p(doc, "Users connect via supported clients; identity via enterprise IdP; pools via Horizon.")
    _h(doc, "Logical"); _p(doc, "Conn Servers, UAG (if external), App Volumes, DEM, Events DB, vCenter integration.")
    _h(doc, "Physical"); _p(doc, f"VCF workload domains: {_safe(ctx,'vcf_domains','(from presales)')}; Regions: {_safe(ctx,'regions','(from presales)')}.")
    doc.save(hld)

    # --- LOE (EPDIO high-level with roles)
    loe = outdir / f"{cust_code}_{proj_code}_LOE_{ts}.docx"
    doc = Document(); _normal_font(doc)
    _h(doc, f"Level of Effort (LOE)\n{_safe(ctx,'project_name')}\n{_safe(ctx,'customer_name')}")
    hours = {"Engage":6,"Plan":8,"Design":14,"Implement":32,"Operate":12}
    _h(doc, "Effort by Phase")
    for k,v in hours.items(): _p(doc, f"{k}: {v} hours")
    _p(doc, f"Total Estimated Hours: {sum(hours.values())}")
    doc.save(loe)

    # --- Bundle
    zip_path = outdir / f"{cust_code}_{proj_code}_Phase1_{ts}.zip"
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as z:
        for p in (sow,hld,loe): z.write(p, arcname=p.name)

    return {"SOW": sow, "HLD": hld, "LOE": loe, "ZIP": zip_path}
